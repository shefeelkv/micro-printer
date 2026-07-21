import io
import json
import logging
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings

from pdf_engine.generator import generate_pdf
from projects.models import CustomFont, Project

logger = logging.getLogger(__name__)

def index(request):
    """Render the primary workspace interface with customized font list."""
    custom_fonts = CustomFont.objects.all()
    # Pre-populate some projects for reference if database is empty
    saved_projects = Project.objects.all()
    
    context = {
        'custom_fonts': custom_fonts,
        'saved_projects': saved_projects,
    }
    return render(request, 'editor/workspace.html', context)

@csrf_exempt
def upload_file(request):
    """
    Endpoint to process uploaded text (.txt) and Word document (.docx) files,
    converting them to structured HTML to seed into the CKEditor.
    """
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Method not allowed.'}, status=405)

    uploaded_file = request.FILES.get('file')
    if not uploaded_file:
        return JsonResponse({'status': 'error', 'message': 'No file uploaded.'}, status=400)

    filename = uploaded_file.name.lower()
    
    try:
        if filename.endswith('.txt'):
            # Simple text parser
            content = uploaded_file.read()
            # Try to decode utf-8, fallback to cp1252
            try:
                text_content = content.decode('utf-8')
            except UnicodeDecodeError:
                text_content = content.decode('cp1252', errors='replace')
            
            # Format lines/paragraphs as HTML
            paragraphs = text_content.strip().split('\n\n')
            paragraphs_cleaned = []
            for p in paragraphs:
                if p.strip():
                    cleaned_text = p.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
                    paragraphs_cleaned.append(f"<p>{cleaned_text}</p>")
            html_content = "".join(paragraphs_cleaned)
            return JsonResponse({'status': 'success', 'html': html_content})

        elif filename.endswith('.docx'):
            # Docx parser
            html_content = docx_to_html(uploaded_file)
            return JsonResponse({'status': 'success', 'html': html_content})

        else:
            return JsonResponse({'status': 'error', 'message': 'Unsupported file type. Please upload a .txt or .docx file.'}, status=400)
            
    except Exception as e:
        logger.exception("Error processing file upload")
        return JsonResponse({'status': 'error', 'message': f"Error parsing file: {str(e)}"}, status=500)

@csrf_exempt
def generate_pdf_view(request):
    """
    Endpoint that processes settings and rich text HTML from client,
    compiles using ReportLab, and returns binary PDF stream for viewer or download.
    """
    if request.method != 'POST':
        return HttpResponse("Method not allowed.", status=405)

    try:
        # Determine source parameters: JSON payload or standard POST
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST

        html_content = data.get('html_content', '')
        # Un-stringify settings if sent as stringified JSON field
        settings_raw = data.get('settings', {})
        if isinstance(settings_raw, str):
            try:
                pdf_settings = json.loads(settings_raw)
            except Exception:
                pdf_settings = {}
        else:
            pdf_settings = settings_raw

        # Run generator
        output_buffer = io.BytesIO()
        generate_pdf(html_content, pdf_settings, output_buffer)
        
        pdf_data = output_buffer.getvalue()
        output_buffer.close()

        response = HttpResponse(pdf_data, content_type='application/pdf')
        
        # Attachment download vs Inline preview
        should_download = request.GET.get('download', 'false').lower() == 'true'
        filename = pdf_settings.get('project_name', 'micro_notes').replace(' ', '_')
        
        if should_download:
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        else:
            response['Content-Disposition'] = f'inline; filename="{filename}.pdf"'
            
        return response

    except Exception as e:
        logger.exception("PDF generation error occurred")
        return HttpResponse(f"PDF generation failed: {str(e)}", status=500)

def docx_to_html(file_stream):
    """Helper to convert docx paragraphs, runs, lists, and tables to formatted HTML."""
    import docx
    
    doc = docx.Document(file_stream)
    html_elements = []
    
    # Track ordered/unordered list elements
    in_list = False
    list_type = None # 'ul' or 'ol'
    
    for p in doc.paragraphs:
        style_name = p.style.name.lower() if p.style else 'normal'
        
        # Parse runs inside the paragraph for inline styles
        p_html = ""
        for run in p.runs:
            text = run.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if not text:
                continue
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            p_html += text
            
        if not p_html.strip():
            continue
            
        # Map paragraph styles
        tag = 'p'
        is_list_item = False
        
        if 'heading 1' in style_name:
            tag = 'h1'
        elif 'heading 2' in style_name:
            tag = 'h2'
        elif 'heading 3' in style_name:
            tag = 'h3'
        elif 'heading 4' in style_name:
            tag = 'h4'
        elif 'heading 5' in style_name:
            tag = 'h5'
        elif 'heading 6' in style_name:
            tag = 'h6'
        elif 'list bullet' in style_name or 'listbullet' in style_name:
            tag = 'li'
            is_list_item = True
            item_list_type = 'ul'
        elif 'list number' in style_name or 'listnumber' in style_name:
            tag = 'li'
            is_list_item = True
            item_list_type = 'ol'
            
        if is_list_item:
            if not in_list:
                in_list = True
                list_type = item_list_type
                html_elements.append(f"<{list_type}>")
            elif list_type != item_list_type:
                html_elements.append(f"</{list_type}>")
                list_type = item_list_type
                html_elements.append(f"<{list_type}>")
            
            html_elements.append(f"<li>{p_html}</li>")
        else:
            if in_list:
                html_elements.append(f"</{list_type}>")
                in_list = False
                list_type = None
            html_elements.append(f"<{tag}>{p_html}</{tag}>")
            
    if in_list:
        html_elements.append(f"</{list_type}>")
        
    # Translate tables
    for table in doc.tables:
        table_html = ['<table border="1" style="width:100%;">']
        for i, row in enumerate(table.rows):
            table_html.append('<tr>')
            for cell in row.cells:
                cell_html = ""
                for cp in cell.paragraphs:
                    cp_html = ""
                    for r in cp.runs:
                        r_text = r.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        if r.bold:
                            r_text = f"<strong>{r_text}</strong>"
                        if r.italic:
                            r_text = f"<em>{r_text}</em>"
                        if r.underline:
                            r_text = f"<u>{r_text}</u>"
                        cp_html += r_text
                    
                    if cp_html.strip():
                        cell_html += f"<p>{cp_html}</p>"
                
                tag = 'th' if i == 0 else 'td'
                cell_html = cell_html or "&nbsp;"
                table_html.append(f"<{tag}>{cell_html}</{tag}>")
            table_html.append('</tr>')
        table_html.append('</table>')
        html_elements.append("\n".join(table_html))
        
    return "\n".join(html_elements)
